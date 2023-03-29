import functools
import os
import xml.etree.ElementTree
import zipfile
from collections import defaultdict, deque
from collections.abc import MutableSequence, Sequence
from pathlib import Path
from typing import MutableSequence, Sequence
from xml.etree.ElementTree import Element
import shutil

import numpy as np
import pandas as pd
import win32com.client as win32
from docx.api import Document
from docx.table import Table

from ..matcher import FieldOrientation
from ..util import table_to_dataframe

_MAX_NESTING_LEVEL = 10


class DocxHandler:
    """Handler for Microsoft Word Documents.

    It is suposed to use the newer .docx format for Word Documents, but if a .doc file is
    provided, it will be converted to .docx format.
    """

    def __init__(self, file_path: Path, temp_folder: Path = Path("temp")) -> None:
        """Creates a new DocxHandler instance.

        Args:
            file_path (Path): Path to file to be handled
            temp_folder (Path, optional): Path to temporary folder where docx files will
                be created when the supplied file has .doc extension. Defaults to 'temp'.
        """
        self._file_path = file_path

        str_file_path = str(file_path.resolve())
        if Path(file_path).suffix[1:] == "doc":
            destination_path = _path_to_docx_file(
                str_file_path, str(temp_folder.resolve())
            )
            _doc_to_docx(
                doc_file_path=str_file_path,
                docx_file_path=destination_path,
            )
            self._file_path = Path(destination_path)

    @property
    def docx_file_path(self) -> str:
        return str(self._file_path.resolve())

    def get_mapping(self, orientation: FieldOrientation) -> dict[str, Sequence[str]]:
        """Retrieves the mapping of values in the document.

        Args:
            orientation (FieldOrientation): Which direction to get the mapping.

        Returns:
            dict[str, Sequence[str]]: A mapping of key names to all contents in the
                document.
        """
        return self._mapping[orientation]

    @functools.cached_property
    def _mapping(self) -> dict[FieldOrientation, dict[str, Sequence[str]]]:
        row_mapping: dict[str, MutableSequence[str]] = defaultdict(list)
        for table in self._docx_tables:
            n_cols = len(table.columns)

            for row in table.rows:
                cells = row.cells
                values = [cell.text for cell in cells]
                for k in range(n_cols - 1):
                    if (k + 1) >= len(values):
                        continue

                    title = values[k]
                    content = values[k + 1]
                    if title == content or title == "" or content == "":
                        continue

                    row_mapping[title].append(content)

        col_mapping: dict[str, MutableSequence[str]] = defaultdict(list)
        for table in self._docx_tables:
            n_rows = len(table.rows)

            for col in table.columns:
                cells = col.cells
                values = [cell.text for cell in cells]
                for k in range(n_rows - 1):
                    if (k + 1) >= len(values):
                        continue

                    title = values[k]
                    content = values[k + 1]
                    if title == content or title == "" or content == "":
                        continue

                    col_mapping[title].append(content)

        return {
            FieldOrientation.ROW: row_mapping,
            FieldOrientation.COLUMN: col_mapping,
        }

    @functools.cache
    def get_tables(self) -> Sequence[pd.DataFrame]:
        """Gets all tables in the document.

        Returns:
            Sequence[pd.DataFrame]: Sequence of tables in the document.
        """
        tables = self._docx_tables

        # Getting subset of tables that has a merged header
        splited_tables = []
        aux_table = []
        for _, table in enumerate(tables):
            for row in table.rows:
                values = [x.text for x in row.cells]
                if len(set(values)) == 1:
                    if len(values[0]) < 200:
                        # this is a merged table header with limited text lenght
                        splited_tables.append(aux_table)
                        aux_table = []
                else:
                    aux_table.append(values)

        dfs = []

        # Gettind data from conventional tables
        for table in tables:
            df = table_to_dataframe(table)

            if len(df) > 0:
                dfs.append(df)

        # Getting data from splited tables
        for table in splited_tables:
            if len(table) == 0:
                continue

            header = table[0]
            data = table[1:]
            header_size = len(header)

            aux = []

            # Removing data with different length from header
            for line in data:
                if len(line) == header_size:
                    aux.append(line)
                else:
                    break

            df = pd.DataFrame(columns=header, data=aux)
            dfs.append(df)

        return _merge_tables(dfs)

    @functools.cached_property
    def _docx_tables(self) -> Sequence[Table]:
        """List of tables in docx document"""
        tables: MutableSequence[Table] = self._document.tables

        new_tables = tables[:]
        count = 0

        # getting all inner tables
        while new_tables:
            aux = []
            for table in new_tables:
                for row in table.rows:
                    for cell in row.cells:
                        aux.extend(cell.tables)

            tables.extend(aux)
            new_tables = aux

            count += 1
            if count > _MAX_NESTING_LEVEL:
                break

        return tables

    @functools.cached_property
    def _document(self) -> Document:
        """Open document and creates a docx file if necessary.

        Returns:
            Document: Word document object.
        """
        return Document(str(self._file_path.resolve()))


_WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_TEXT = _WORD_NAMESPACE + "t"
_TABLE = _WORD_NAMESPACE + "tbl"
_ROW = _WORD_NAMESPACE + "tr"
_CELL = _WORD_NAMESPACE + "tc"
_PARAGRAPH = _WORD_NAMESPACE + "p"


class DocxXMLHandler:
    def __init__(
        self, file_path: str, temp_folder: str = str(Path("temp").resolve())
    ) -> None:
        """Creates a new DocxXMLHandler instance.

        This handler does not rely on Word to get information from the files, resulting in
        faster extraction and avoid the program to get stuck in an API call.

        Args:
            file_path (str): File path to be handled.
            temp_folder (str, optional): Path to temporary folder where docx files will
                be created when the supplied file has .doc extension. Defaults to 'temp'.
        """
        self._file_path = file_path

        if Path(file_path).suffix[1:] == "doc":
            destination_path = _path_to_docx_file(file_path, temp_folder)
            _doc_to_docx(
                doc_file_path=file_path,
                docx_file_path=destination_path,
            )
            self._file_path = destination_path

    @property
    def docx_file_path(self) -> str:
        return self._file_path

    @functools.cache
    def get_tables(self) -> Sequence[pd.DataFrame]:
        with zipfile.ZipFile(self._file_path) as docx:
            tree = xml.etree.ElementTree.XML(docx.read("word/document.xml"))

        tables: MutableSequence[pd.DataFrame] = []
        for table in tree.iter(_TABLE):
            data: MutableSequence[MutableSequence[str]] = []
            for row in table.findall(f"./{_ROW}"):
                data.append(
                    [_get_combined_text(cell) for cell in row.findall(f"./{_CELL}")]
                )

            df = pd.DataFrame(data)
            df.columns = df.iloc[0].to_numpy(dtype=str)
            df = df[1:]

            tables.append(df)

        return _merge_tables(tables)

    @functools.cached_property
    def _mapping(self) -> dict[FieldOrientation, dict[str, Sequence[str]]]:
        mapping: dict[FieldOrientation, dict[str, MutableSequence[str]]] = {
            FieldOrientation.ROW: defaultdict(list),
            FieldOrientation.COLUMN: defaultdict(list),
        }
        for df in self.get_tables():
            records = _table_to_records(df)
            for row in records:
                for i, key in enumerate(row[:-1]):
                    mapping[FieldOrientation.ROW][key].append(row[i + 1])

            for col in np.array(records).T.tolist():
                for i, key in enumerate(col[:-1]):
                    mapping[FieldOrientation.COLUMN][key].append(col[i + 1])

        return mapping

    def get_mapping(self, orientation: FieldOrientation) -> dict[str, Sequence[str]]:
        """Retrieves the mapping of values in the document.

        Args:
            orientation (FieldOrientation): Which direction to get the mapping.

        Returns:
            dict[str, Sequence[str]]: A mapping of key names to all contents in the
                document.
        """
        return self._mapping[orientation]


def _table_to_records(df: pd.DataFrame) -> Sequence[Sequence[str]]:
    return [df.columns.to_numpy(dtype=str).tolist()] + df.to_numpy(dtype=str).tolist()


def _get_combined_text(cell: Element) -> str:
    # TODO handle merged cells
    fragments: MutableSequence[str] = []
    q = deque(cell.findall("./*"))
    while q:
        cell = q.popleft()
        if cell.tag == _TABLE:
            continue

        if cell.tag == _PARAGRAPH:
            if fragments:
                fragments.append("\n")
        elif cell.tag == _TEXT:
            if frag := str(cell.text):
                fragments.append(frag)

        q.extendleft(cell.findall("./*")[::-1])

    return "".join(fragments)


def _path_to_docx_file(doc_file_path: str, folder: str) -> str:
    """Creates a standard name for temporary .docx files."""
    folder_path = Path(folder)
    folder_path.mkdir(parents=True, exist_ok=True)

    original_file_name = Path(doc_file_path).stem
    destination_path = folder_path / f"x_{original_file_name}.docx"

    return str(destination_path.resolve())


def _merge_tables(tables: Sequence[pd.DataFrame]) -> Sequence[pd.DataFrame]:
    groups: dict[str, MutableSequence[pd.DataFrame]] = defaultdict(list)
    for df in tables:
        cols = df.columns.to_numpy(dtype=str).tolist()
        cols.sort()
        groups["".join(cols)].append(df)

    merged_tables: Sequence[pd.DataFrame] = []
    for dfs in groups.values():
        merged = pd.concat(dfs)
        merged.drop_duplicates(inplace=True)
        merged.reset_index(drop=True, inplace=True)
        merged_tables.append(merged)

    return merged_tables


def _doc_to_docx(doc_file_path: str, docx_file_path: str) -> None:
    """Converts a .doc file to a .docx file.

    If a document in the docx_file_path already exists, this function does nothing.

    Args:
        doc_file_path (str): Path to the .doc file that will be converted.
        docx_file_path (str): Path where the new file will be created.
    """
    if os.path.exists(docx_file_path):
        return

    try:
        shutil.rmtree(_get_python_temp_appdata_folder())
    except Exception:
        pass

    word = win32.gencache.EnsureDispatch("Word.Application")

    doc = word.Documents.Open(doc_file_path)
    doc.Activate()

    word.ActiveDocument.SaveAs(
        docx_file_path,
        FileFormat=win32.constants.wdFormatXMLDocument,
    )
    doc.Close(False)


def _get_python_temp_appdata_folder() -> str:
    return f'{os.getenv("LOCALAPPDATA")}\Local\Temp'
